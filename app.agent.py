import os
import json
import base64
import sqlite3
from collections import Counter
from datetime import datetime
from pathlib import Path

import streamlit as st
from dotenv import load_dotenv
from docx import Document
from openai import OpenAI
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail,
    Attachment,
    FileContent,
    FileName,
    FileType,
    Disposition,
)

load_dotenv()

DATA_DIR = Path("data")
OUTPUT_DIR = Path("outputs")

DATA_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

DB_PATH = DATA_DIR / "study_agent_memory.db"


# -----------------------------
# Database setup
# -----------------------------
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS student_profiles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_name TEXT NOT NULL,
            year_group TEXT NOT NULL,
            tutor_email TEXT NOT NULL,
            created_at TEXT,
            updated_at TEXT,
            UNIQUE(student_name, tutor_email)
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS chat_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_name TEXT,
            role TEXT,
            message TEXT,
            created_at TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS learning_notes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_name TEXT,
            year_group TEXT,
            subject TEXT,
            topic TEXT,
            note TEXT,
            created_at TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS student_struggles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_name TEXT,
            subject TEXT,
            topic TEXT,
            struggle_area TEXT,
            evidence TEXT,
            suggested_support TEXT,
            created_at TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS quiz_attempts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_name TEXT,
            subject TEXT,
            topic TEXT,
            question TEXT,
            student_answer TEXT,
            correct_answer TEXT,
            was_correct TEXT,
            feedback TEXT,
            created_at TEXT
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS active_quiz_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_name TEXT,
            subject TEXT,
            topic TEXT,
            question TEXT,
            expected_answer TEXT,
            created_at TEXT
        )
        """
    )

    conn.commit()
    conn.close()


# -----------------------------
# Helpers
# -----------------------------
def current_timestamp():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def clean_text(value):
    return value.strip() if value else ""


def clean_email(value):
    return value.strip().lower() if value else ""


# -----------------------------
# Student profile memory
# -----------------------------
def save_or_update_student_profile(student_name, year_group, tutor_email):
    student_name = clean_text(student_name)
    year_group = clean_text(year_group)
    tutor_email = clean_email(tutor_email)
    now = current_timestamp()

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        """
        INSERT INTO student_profiles (
            student_name,
            year_group,
            tutor_email,
            created_at,
            updated_at
        )
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(student_name, tutor_email)
        DO UPDATE SET
            year_group = excluded.year_group,
            updated_at = excluded.updated_at
        """,
        (student_name, year_group, tutor_email, now, now),
    )

    conn.commit()
    conn.close()

    return {
        "status": "saved",
        "student_name": student_name,
        "year_group": year_group,
        "tutor_email": tutor_email,
    }


def get_students_for_tutor(tutor_email):
    tutor_email = clean_email(tutor_email)

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT student_name, year_group, tutor_email
        FROM student_profiles
        WHERE lower(tutor_email) = lower(?)
        ORDER BY student_name
        """,
        (tutor_email,),
    )

    rows = cursor.fetchall()
    conn.close()

    students = []
    for student_name, year_group, saved_tutor_email in rows:
        students.append(
            {
                "student_name": student_name,
                "year_group": year_group,
                "tutor_email": saved_tutor_email,
            }
        )

    return students


# -----------------------------
# Chat memory
# -----------------------------
def save_chat_message(student_name, role, message):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        """
        INSERT INTO chat_messages (student_name, role, message, created_at)
        VALUES (?, ?, ?, ?)
        """,
        (student_name, role, message, current_timestamp()),
    )

    conn.commit()
    conn.close()


def get_recent_chat(student_name, limit=12):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT role, message
        FROM chat_messages
        WHERE student_name = ?
        ORDER BY id DESC
        LIMIT ?
        """,
        (student_name, limit),
    )

    rows = cursor.fetchall()
    conn.close()

    rows.reverse()
    return rows


# -----------------------------
# Learning note memory
# -----------------------------
def save_learning_note(student_name, year_group, subject, topic, note):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        """
        INSERT INTO learning_notes (
            student_name,
            year_group,
            subject,
            topic,
            note,
            created_at
        )
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            student_name,
            year_group,
            subject,
            topic,
            note,
            current_timestamp(),
        ),
    )

    conn.commit()
    conn.close()

    return {
        "status": "saved",
        "student_name": student_name,
        "subject": subject,
        "topic": topic,
        "message": "Learning note saved successfully.",
    }


def get_previous_notes(student_name, subject=None, topic=None, limit=10):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    subject = clean_text(subject)
    topic = clean_text(topic)

    if subject and topic:
        cursor.execute(
            """
            SELECT subject, topic, note, created_at
            FROM learning_notes
            WHERE student_name = ?
              AND lower(subject) LIKE lower(?)
              AND lower(topic) LIKE lower(?)
            ORDER BY id DESC
            LIMIT ?
            """,
            (student_name, f"%{subject}%", f"%{topic}%", limit),
        )
    elif subject:
        cursor.execute(
            """
            SELECT subject, topic, note, created_at
            FROM learning_notes
            WHERE student_name = ?
              AND lower(subject) LIKE lower(?)
            ORDER BY id DESC
            LIMIT ?
            """,
            (student_name, f"%{subject}%", limit),
        )
    else:
        cursor.execute(
            """
            SELECT subject, topic, note, created_at
            FROM learning_notes
            WHERE student_name = ?
            ORDER BY id DESC
            LIMIT ?
            """,
            (student_name, limit),
        )

    rows = cursor.fetchall()
    conn.close()

    notes = []
    for saved_subject, saved_topic, note, created_at in rows:
        notes.append(
            {
                "subject": saved_subject,
                "topic": saved_topic,
                "note": note,
                "created_at": created_at,
            }
        )

    return notes


# -----------------------------
# Active quiz session memory
# -----------------------------
def save_active_quiz_question(student_name, subject, topic, question, expected_answer):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    # Keep only one active quiz question per student/subject/topic.
    cursor.execute(
        """
        DELETE FROM active_quiz_sessions
        WHERE student_name = ?
          AND lower(subject) = lower(?)
          AND lower(topic) = lower(?)
        """,
        (student_name, subject, topic),
    )

    cursor.execute(
        """
        INSERT INTO active_quiz_sessions (
            student_name,
            subject,
            topic,
            question,
            expected_answer,
            created_at
        )
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            student_name,
            subject,
            topic,
            question,
            expected_answer,
            current_timestamp(),
        ),
    )

    conn.commit()
    conn.close()

    return {
        "status": "saved",
        "student_name": student_name,
        "subject": subject,
        "topic": topic,
        "question": question,
        "expected_answer": expected_answer,
        "message": "Active quiz question saved successfully.",
    }


def get_active_quiz_question(student_name, subject=None, topic=None):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    subject = clean_text(subject)
    topic = clean_text(topic)

    if subject and topic:
        cursor.execute(
            """
            SELECT subject, topic, question, expected_answer, created_at
            FROM active_quiz_sessions
            WHERE student_name = ?
              AND lower(subject) = lower(?)
              AND lower(topic) = lower(?)
            ORDER BY id DESC
            LIMIT 1
            """,
            (student_name, subject, topic),
        )
    else:
        cursor.execute(
            """
            SELECT subject, topic, question, expected_answer, created_at
            FROM active_quiz_sessions
            WHERE student_name = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (student_name,),
        )

    row = cursor.fetchone()
    conn.close()

    if not row:
        return {
            "status": "not_found",
            "message": "No active quiz question found for this student.",
        }

    saved_subject, saved_topic, question, expected_answer, created_at = row

    return {
        "status": "found",
        "student_name": student_name,
        "subject": saved_subject,
        "topic": saved_topic,
        "question": question,
        "expected_answer": expected_answer,
        "created_at": created_at,
    }


def clear_active_quiz_question(student_name, subject=None, topic=None):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    subject = clean_text(subject)
    topic = clean_text(topic)

    if subject and topic:
        cursor.execute(
            """
            DELETE FROM active_quiz_sessions
            WHERE student_name = ?
              AND lower(subject) = lower(?)
              AND lower(topic) = lower(?)
            """,
            (student_name, subject, topic),
        )
    else:
        cursor.execute(
            """
            DELETE FROM active_quiz_sessions
            WHERE student_name = ?
            """,
            (student_name,),
        )

    deleted_count = cursor.rowcount
    conn.commit()
    conn.close()

    return {
        "status": "cleared",
        "deleted_count": deleted_count,
        "message": "Active quiz question cleared.",
    }


# -----------------------------
# Struggle memory
# -----------------------------
def record_struggle(student_name, subject, topic, struggle_area, evidence, suggested_support):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        """
        INSERT INTO student_struggles (
            student_name,
            subject,
            topic,
            struggle_area,
            evidence,
            suggested_support,
            created_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            student_name,
            subject,
            topic,
            struggle_area,
            evidence,
            suggested_support,
            current_timestamp(),
        ),
    )

    conn.commit()
    conn.close()

    return {
        "status": "recorded",
        "student_name": student_name,
        "subject": subject,
        "topic": topic,
        "struggle_area": struggle_area,
        "message": "Student struggle area recorded successfully.",
    }


def get_struggles(student_name):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT subject, topic, struggle_area, evidence, suggested_support, created_at
        FROM student_struggles
        WHERE student_name = ?
        ORDER BY id DESC
        """,
        (student_name,),
    )

    rows = cursor.fetchall()
    conn.close()

    struggles = []
    for subject, topic, area, evidence, support, created_at in rows:
        struggles.append(
            {
                "subject": subject,
                "topic": topic,
                "struggle_area": area,
                "evidence": evidence,
                "suggested_support": support,
                "created_at": created_at,
            }
        )

    return struggles


# -----------------------------
# Quiz memory
# -----------------------------
def record_quiz_attempt(
    student_name,
    subject,
    topic,
    question,
    student_answer,
    correct_answer,
    was_correct,
    feedback,
):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute(
        """
        INSERT INTO quiz_attempts (
            student_name,
            subject,
            topic,
            question,
            student_answer,
            correct_answer,
            was_correct,
            feedback,
            created_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            student_name,
            subject,
            topic,
            question,
            student_answer,
            correct_answer,
            was_correct,
            feedback,
            current_timestamp(),
        ),
    )

    conn.commit()
    conn.close()

    return {
        "status": "recorded",
        "student_name": student_name,
        "subject": subject,
        "topic": topic,
        "was_correct": was_correct,
        "message": "Quiz attempt recorded successfully.",
    }


def get_quiz_attempts(student_name, subject=None, topic=None, limit=30):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    subject = clean_text(subject)
    topic = clean_text(topic)

    if subject and topic:
        cursor.execute(
            """
            SELECT subject, topic, question, student_answer, correct_answer, was_correct, feedback, created_at
            FROM quiz_attempts
            WHERE student_name = ?
              AND lower(subject) LIKE lower(?)
              AND lower(topic) LIKE lower(?)
            ORDER BY id DESC
            LIMIT ?
            """,
            (student_name, f"%{subject}%", f"%{topic}%", limit),
        )
    elif subject:
        cursor.execute(
            """
            SELECT subject, topic, question, student_answer, correct_answer, was_correct, feedback, created_at
            FROM quiz_attempts
            WHERE student_name = ?
              AND lower(subject) LIKE lower(?)
            ORDER BY id DESC
            LIMIT ?
            """,
            (student_name, f"%{subject}%", limit),
        )
    else:
        cursor.execute(
            """
            SELECT subject, topic, question, student_answer, correct_answer, was_correct, feedback, created_at
            FROM quiz_attempts
            WHERE student_name = ?
            ORDER BY id DESC
            LIMIT ?
            """,
            (student_name, limit),
        )

    rows = cursor.fetchall()
    conn.close()

    attempts = []
    for subject, topic, question, student_answer, correct_answer, was_correct, feedback, created_at in rows:
        attempts.append(
            {
                "subject": subject,
                "topic": topic,
                "question": question,
                "student_answer": student_answer,
                "correct_answer": correct_answer,
                "was_correct": was_correct,
                "feedback": feedback,
                "created_at": created_at,
            }
        )

    return attempts


# -----------------------------
# Report/document/email tools
# -----------------------------
def create_dashboard_summary(student_name):
    notes = get_previous_notes(student_name, limit=100)
    quiz_attempts = get_quiz_attempts(student_name, limit=100)
    struggles = get_struggles(student_name)

    total_notes = len(notes)
    total_quiz_attempts = len(quiz_attempts)

    correct_count = 0
    partial_count = 0
    incorrect_count = 0

    for attempt in quiz_attempts:
        result = clean_text(attempt.get("was_correct", "")).lower()

        if result in ["yes", "correct", "true"]:
            correct_count += 1
        elif result in ["partial", "partially correct", "partly correct"]:
            partial_count += 1
        else:
            incorrect_count += 1

    latest_activity = "No activity yet"
    activity_dates = []

    for item in notes:
        activity_dates.append(item.get("created_at", ""))

    for item in quiz_attempts:
        activity_dates.append(item.get("created_at", ""))

    for item in struggles:
        activity_dates.append(item.get("created_at", ""))

    valid_dates = [date for date in activity_dates if date]

    if valid_dates:
        latest_activity = max(valid_dates)

    weak_topic_counter = Counter()

    for item in struggles:
        subject = item.get("subject", "Unknown subject") or "Unknown subject"
        topic = item.get("topic", "Unknown topic") or "Unknown topic"
        weak_topic_counter[f"{subject}: {topic}"] += 1

    weakest_topics = weak_topic_counter.most_common(5)

    accuracy = 0
    if total_quiz_attempts > 0:
        accuracy = round((correct_count / total_quiz_attempts) * 100, 1)

    return {
        "total_notes": total_notes,
        "total_quiz_attempts": total_quiz_attempts,
        "correct_count": correct_count,
        "partial_count": partial_count,
        "incorrect_count": incorrect_count,
        "accuracy": accuracy,
        "latest_activity": latest_activity,
        "weakest_topics": weakest_topics,
    }


# -----------------------------
# Report/document/email tools
# -----------------------------
def create_parent_tutor_report_from_memory(student_name):
    notes = get_previous_notes(student_name, limit=30)
    struggles = get_struggles(student_name)
    quiz_attempts = get_quiz_attempts(student_name, limit=50)
    active_quiz = get_active_quiz_question(student_name)

    return {
        "student_name": student_name,
        "learning_notes": notes,
        "struggles": struggles,
        "quiz_attempts": quiz_attempts,
        "active_quiz": active_quiz,
    }


def create_revision_document(student_name, subject, topic, document_title, document_content):
    safe_student = student_name.replace(" ", "_").lower() or "student"
    safe_subject = subject.replace(" ", "_").lower() if subject else "general"
    safe_topic = topic.replace(" ", "_").lower() if topic else "revision"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    filename = f"{safe_student}_{safe_subject}_{safe_topic}_{timestamp}.docx"
    file_path = OUTPUT_DIR / filename

    doc = Document()

    doc.add_heading(document_title, level=1)
    doc.add_paragraph(f"Student: {student_name}")

    if subject:
        doc.add_paragraph(f"Subject: {subject}")

    if topic:
        doc.add_paragraph(f"Topic: {topic}")

    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    doc.add_heading("Study Content", level=2)

    for paragraph in document_content.split("\n"):
        paragraph = paragraph.strip()
        if paragraph:
            doc.add_paragraph(paragraph)

    doc.save(file_path)

    return {
        "status": "created",
        "file_path": str(file_path),
        "file_name": filename,
        "message": "Revision document created successfully.",
    }


def send_email_with_attachment(to_email, email_subject, email_body, attachment_path):
    sendgrid_api_key = get_secret("SENDGRID_API_KEY")
    from_email = get_secret("FROM_EMAIL")

    if not sendgrid_api_key:
        return {
            "status": "failed",
            "message": "SENDGRID_API_KEY is missing in the .env file.",
        }

    if not from_email:
        return {
            "status": "failed",
            "message": "FROM_EMAIL is missing in the .env file.",
        }

    file_path = Path(attachment_path)

    if not file_path.exists():
        return {
            "status": "failed",
            "message": f"Attachment file not found: {attachment_path}",
        }

    message = Mail(
        from_email=from_email,
        to_emails=to_email,
        subject=email_subject,
        plain_text_content=email_body,
    )

    with open(file_path, "rb") as file:
        encoded_file = base64.b64encode(file.read()).decode()

    attachment = Attachment(
        FileContent(encoded_file),
        FileName(file_path.name),
        FileType("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        Disposition("attachment"),
    )

    message.attachment = attachment

    sg = SendGridAPIClient(sendgrid_api_key)
    response = sg.send(message)

    return {
        "status": "sent",
        "status_code": response.status_code,
        "to_email": to_email,
        "file_name": file_path.name,
        "message": "Email sent successfully.",
    }


# -----------------------------
# OpenAI setup
# -----------------------------

def get_secret(name, default=None):
    try:
        return st.secrets[name]
    except Exception:
        return os.getenv(name, default)

def get_openai_client():
    api_key = get_secret("OPENAI_API_KEY")

    if not api_key:
        return None

    return OpenAI(api_key=api_key)


# -----------------------------
# Agent tool definitions
# -----------------------------
AGENT_TOOLS = [
    {
        "type": "function",
        "name": "save_learning_note",
        "description": "Save a student's learning note from today's lesson into long-term memory.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
                "year_group": {"type": "string"},
                "subject": {"type": "string"},
                "topic": {"type": "string"},
                "note": {"type": "string"},
            },
            "required": ["student_name", "year_group", "subject", "topic", "note"],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "get_previous_notes",
        "description": "Retrieve previous notes from memory for a student, optionally filtered by subject and topic.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
                "subject": {"type": "string"},
                "topic": {"type": "string"},
                "limit": {"type": "integer"},
            },
            "required": ["student_name"],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "save_active_quiz_question",
        "description": "Save the current active quiz question and expected answer before asking the student to answer it.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
                "subject": {"type": "string"},
                "topic": {"type": "string"},
                "question": {"type": "string"},
                "expected_answer": {"type": "string"},
            },
            "required": ["student_name", "subject", "topic", "question", "expected_answer"],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "get_active_quiz_question",
        "description": "Retrieve the active quiz question the student is expected to answer.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
                "subject": {"type": "string"},
                "topic": {"type": "string"},
            },
            "required": ["student_name"],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "clear_active_quiz_question",
        "description": "Clear the active quiz question after the student's answer has been checked and recorded.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
                "subject": {"type": "string"},
                "topic": {"type": "string"},
            },
            "required": ["student_name"],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "record_struggle",
        "description": "Record a student's misunderstanding, weak area, incorrect answer, or topic needing support.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
                "subject": {"type": "string"},
                "topic": {"type": "string"},
                "struggle_area": {"type": "string"},
                "evidence": {"type": "string"},
                "suggested_support": {"type": "string"},
            },
            "required": [
                "student_name",
                "subject",
                "topic",
                "struggle_area",
                "evidence",
                "suggested_support",
            ],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "record_quiz_attempt",
        "description": "Record a student's answer to a quiz question, including whether it was correct and the feedback given.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
                "subject": {"type": "string"},
                "topic": {"type": "string"},
                "question": {"type": "string"},
                "student_answer": {"type": "string"},
                "correct_answer": {"type": "string"},
                "was_correct": {"type": "string"},
                "feedback": {"type": "string"},
            },
            "required": [
                "student_name",
                "subject",
                "topic",
                "question",
                "student_answer",
                "correct_answer",
                "was_correct",
                "feedback",
            ],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "get_quiz_attempts",
        "description": "Retrieve previous quiz attempts for a student, optionally filtered by subject and topic.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
                "subject": {"type": "string"},
                "topic": {"type": "string"},
                "limit": {"type": "integer"},
            },
            "required": ["student_name"],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "create_parent_tutor_report_from_memory",
        "description": "Retrieve memory needed to create a parent/tutor report about the student's progress and struggles.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
            },
            "required": ["student_name"],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "create_revision_document",
        "description": "Create a Word document for a student's revision notes, quiz summary, study pack, or parent/tutor report.",
        "parameters": {
            "type": "object",
            "properties": {
                "student_name": {"type": "string"},
                "subject": {"type": "string"},
                "topic": {"type": "string"},
                "document_title": {"type": "string"},
                "document_content": {"type": "string"},
            },
            "required": [
                "student_name",
                "subject",
                "topic",
                "document_title",
                "document_content",
            ],
            "additionalProperties": False,
        },
    },
    {
        "type": "function",
        "name": "send_email_with_attachment",
        "description": "Send an email with a Word document attachment. Only use when the user clearly asks to email a document and provides a recipient email address.",
        "parameters": {
            "type": "object",
            "properties": {
                "to_email": {"type": "string"},
                "email_subject": {"type": "string"},
                "email_body": {"type": "string"},
                "attachment_path": {"type": "string"},
            },
            "required": ["to_email", "email_subject", "email_body", "attachment_path"],
            "additionalProperties": False,
        },
    },
]


def execute_tool(tool_name, arguments):
    if tool_name == "save_learning_note":
        return save_learning_note(**arguments)
    if tool_name == "get_previous_notes":
        return get_previous_notes(**arguments)
    if tool_name == "save_active_quiz_question":
        return save_active_quiz_question(**arguments)
    if tool_name == "get_active_quiz_question":
        return get_active_quiz_question(**arguments)
    if tool_name == "clear_active_quiz_question":
        return clear_active_quiz_question(**arguments)
    if tool_name == "record_struggle":
        return record_struggle(**arguments)
    if tool_name == "record_quiz_attempt":
        return record_quiz_attempt(**arguments)
    if tool_name == "get_quiz_attempts":
        return get_quiz_attempts(**arguments)
    if tool_name == "create_parent_tutor_report_from_memory":
        return create_parent_tutor_report_from_memory(**arguments)
    if tool_name == "create_revision_document":
        return create_revision_document(**arguments)
    if tool_name == "send_email_with_attachment":
        return send_email_with_attachment(**arguments)

    return {"error": f"Unknown tool: {tool_name}"}


# -----------------------------
# Agent loop
# -----------------------------
def run_study_agent(student_name, year_group, subject, topic, user_message):
    client = get_openai_client()

    if client is None:
        return "OpenAI API key is missing. Please add OPENAI_API_KEY to your .env file."

    model = get_secret("OPENAI_MODEL", "gpt-5.2")

    recent_chat = get_recent_chat(student_name, limit=10)
    recent_chat_text = ""

    for role, message in recent_chat:
        recent_chat_text += f"{role}: {message}\n"

    system_prompt = f"""
You are an agentic AI study tutor for a school student.

Student context:
- Student name: {student_name}
- Year/Class: {year_group}
- Current subject: {subject}
- Current topic: {topic}

Recent chat:
{recent_chat_text}

You have tools available.

Use tools when helpful:
- Use save_learning_note when the student shares what they learnt today.
- Use get_previous_notes when the student asks to revise, asks for a quiz, or asks about something learnt earlier.
- When the student asks to be quizzed, use previous notes if available, create one question, call save_active_quiz_question with the question and expected answer, then ask the student that one question.
- When the student answers a quiz question, call get_active_quiz_question first, compare the answer with the saved expected answer, then call record_quiz_attempt.
- After recording a quiz attempt, call clear_active_quiz_question.
- If the student answer is wrong or partially correct, call record_struggle as well.
- Use get_quiz_attempts when the parent/tutor asks about quiz performance, wrong answers, improvement, or weak areas.
- Use create_parent_tutor_report_from_memory when the parent/tutor asks for a progress or struggle report.
- Use create_revision_document when the student or parent/tutor asks to create a Word document, study pack, revision notes, quiz summary, or parent/tutor report document.
- Use send_email_with_attachment only when the user clearly asks to email a document, provides an email address, and there is a document path available from a previous document creation tool result.

Important behaviour:
- Ask quiz questions one at a time.
- Do not ask multiple quiz questions in one response.
- When correcting an answer, be kind and specific.
- Do not invent previous learning history. Use memory tools where needed.
- Keep the tone supportive, clear, age-appropriate and concise.
- For learning notes, summarise them and confirm they have been saved.
- For parent/tutor reports, use memory and clearly separate evidence from suggestions.
- Do not send an email unless the user has clearly asked for it.
- If the user has not provided an email address or document path for email, ask for the missing information.
"""

    response = client.responses.create(
        model=model,
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        tools=AGENT_TOOLS,
    )

    # Allow the model to call more than one tool in sequence.
    for _ in range(7):
        tool_outputs = []

        for item in response.output:
            if item.type == "function_call":
                try:
                    tool_args = json.loads(item.arguments)
                    tool_result = execute_tool(item.name, tool_args)
                except Exception as error:
                    tool_result = {"status": "failed", "message": str(error)}

                tool_outputs.append(
                    {
                        "type": "function_call_output",
                        "call_id": item.call_id,
                        "output": json.dumps(tool_result),
                    }
                )

        if not tool_outputs:
            return response.output_text

        response = client.responses.create(
            model=model,
            input=tool_outputs,
            previous_response_id=response.id,
            tools=AGENT_TOOLS,
        )

    return response.output_text or "The agent completed tool actions but did not produce a final message."


# -----------------------------
# Streamlit interface
# -----------------------------
st.set_page_config(
    page_title="Agentic AI Study Tutor",
    page_icon="🤖",
    layout="wide",
)

init_db()

st.title("🤖 Agentic AI Study Tutor")
st.write(
    "This prototype lets the AI tutor save learning notes, retrieve previous learning, manage active quiz sessions, record quiz attempts, track struggles, create reports, and create documents."
)

# -----------------------------
# Login
# -----------------------------
st.sidebar.subheader("Login")

role = st.sidebar.selectbox(
    "Select role",
    ["Student", "Parent/Tutor"],
)

passcode = st.sidebar.text_input("Passcode", type="password")

STUDENT_PASSCODE = get_secret("STUDENT_PASSCODE", "student123")
TUTOR_PASSCODE = get_secret("TUTOR_PASSCODE", "tutor123")

if role == "Student" and passcode != STUDENT_PASSCODE:
    st.warning("Please enter the student passcode.")
    st.stop()

if role == "Parent/Tutor" and passcode != TUTOR_PASSCODE:
    st.warning("Please enter the parent/tutor passcode.")
    st.stop()

# -----------------------------
# Created documents sidebar
# -----------------------------
st.sidebar.subheader("Created Documents")

created_files = list(OUTPUT_DIR.glob("*.docx"))

if not created_files:
    st.sidebar.info("No Word documents created yet.")
else:
    latest_files = sorted(
        created_files,
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )[:5]

    for file_path in latest_files:
        with open(file_path, "rb") as file:
            st.sidebar.download_button(
                label=f"Download {file_path.name}",
                data=file,
                file_name=file_path.name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{file_path.name}",
            )

# -----------------------------
# Student view
# -----------------------------
if role == "Student":
    st.subheader("Student Agent Chat")

    col1, col2 = st.columns(2)

    with col1:
        student_name = st.text_input("Student name", placeholder="Enter student name")
        year_group = st.text_input("Year/Class", placeholder="Example: Year 8")

    with col2:
        subject = st.text_input("Subject", placeholder="Example: Biology")
        topic = st.text_input("Topic", placeholder="Example: Photosynthesis")

    tutor_email = st.text_input(
        "Parent/Tutor email to link this student",
        placeholder="Example: tutor@example.com",
    )

    if not student_name or not year_group or not subject or not topic or not tutor_email:
        st.info("Please enter student name, year/class, subject, topic, and parent/tutor email to start.")
        st.stop()

    save_or_update_student_profile(
        student_name=student_name,
        year_group=year_group,
        tutor_email=tutor_email,
    )

    st.success(f"Student linked to parent/tutor email: {tutor_email}")
    st.info("Type today's notes, ask for a quiz, answer a question, or ask for a revision document.")

    previous_chat = get_recent_chat(student_name, limit=20)

    for chat_role, message in previous_chat:
        with st.chat_message(chat_role):
            st.write(message)

    active_quiz = get_active_quiz_question(student_name, subject=subject, topic=topic)

    if active_quiz.get("status") == "found":
        st.warning(f"Active quiz question: {active_quiz['question']}")

    user_message = st.chat_input("Type your notes, ask for a quiz, or answer a question")

    if user_message:
        save_chat_message(student_name, "user", user_message)

        with st.chat_message("user"):
            st.write(user_message)

        with st.spinner("Agent is thinking and using tools if needed..."):
            assistant_reply = run_study_agent(
                student_name=student_name,
                year_group=year_group,
                subject=subject,
                topic=topic,
                user_message=user_message,
            )

        save_chat_message(student_name, "assistant", assistant_reply)

        with st.chat_message("assistant"):
            st.write(assistant_reply)

# -----------------------------
# Parent/Tutor view
# -----------------------------
if role == "Parent/Tutor":
    tab_tutor, tab_memory = st.tabs(["Parent/Tutor Agent", "Memory Viewer"])

    with tab_tutor:
        st.subheader("Parent/Tutor Agent")

        tutor_email = st.text_input(
            "Parent/Tutor email",
            placeholder="Example: tutor@example.com",
            key="tutor_email_report",
        )

        if not tutor_email:
            st.info("Enter the parent/tutor email to view linked students.")
            st.stop()

        linked_students = get_students_for_tutor(tutor_email)

        if not linked_students:
            st.warning("No students are currently linked to this parent/tutor email.")
            st.stop()

        student_options = [
            f"{student['student_name']} — {student['year_group']}"
            for student in linked_students
        ]

        selected_student_label = st.selectbox(
            "Select linked student",
            student_options,
            key="report_student_select",
        )

        selected_student = linked_students[student_options.index(selected_student_label)]
        selected_student_name = selected_student["student_name"]
        selected_year_group = selected_student["year_group"]

        tutor_question = st.text_area(
            "Ask about the student's progress or struggles",
            value="Where is this student struggling and what should we focus on?",
            height=120,
        )

        if st.button("Ask Parent/Tutor Agent"):
            tutor_prompt = f"""
Parent/tutor question for student {selected_student_name}:
{tutor_question}

Please create a professional parent/tutor response using memory.
"""

            with st.spinner("Parent/tutor agent is checking memory..."):
                report = run_study_agent(
                    student_name=selected_student_name,
                    year_group=selected_year_group,
                    subject="",
                    topic="",
                    user_message=tutor_prompt,
                )

            st.subheader("Agent Report")
            st.write(report)

    with tab_memory:
        st.subheader("Memory Viewer")

        memory_tutor_email = st.text_input(
            "Parent/Tutor email",
            placeholder="Example: tutor@example.com",
            key="tutor_email_memory",
        )

        if not memory_tutor_email:
            st.info("Enter the parent/tutor email to view linked student memory.")
            st.stop()

        memory_linked_students = get_students_for_tutor(memory_tutor_email)

        if not memory_linked_students:
            st.warning("No students are currently linked to this parent/tutor email.")
            st.stop()

        memory_student_options = [
            f"{student['student_name']} — {student['year_group']}"
            for student in memory_linked_students
        ]

        selected_memory_student_label = st.selectbox(
            "Select linked student",
            memory_student_options,
            key="memory_student_select",
        )

        selected_memory_student = memory_linked_students[
            memory_student_options.index(selected_memory_student_label)
        ]
        memory_student_name = selected_memory_student["student_name"]

        st.markdown("### Progress Dashboard")
        dashboard = create_dashboard_summary(memory_student_name)

        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.metric("Learning notes", dashboard["total_notes"])

        with col2:
            st.metric("Quiz attempts", dashboard["total_quiz_attempts"])

        with col3:
            st.metric("Correct answers", dashboard["correct_count"])

        with col4:
            st.metric("Accuracy", f"{dashboard['accuracy']}%")

        col5, col6, col7 = st.columns(3)

        with col5:
            st.metric("Partially correct", dashboard["partial_count"])

        with col6:
            st.metric("Incorrect", dashboard["incorrect_count"])

        with col7:
            st.metric("Latest activity", dashboard["latest_activity"])

        st.markdown("#### Common Weak Topics")

        if not dashboard["weakest_topics"]:
            st.info("No weak topics recorded yet.")
        else:
            for topic_label, count in dashboard["weakest_topics"]:
                st.write(f"- **{topic_label}** — {count} recorded issue(s)")

        st.markdown("### Active Quiz Question")
        active_quiz = get_active_quiz_question(memory_student_name)

        if active_quiz.get("status") == "found":
            st.write(f"**Subject:** {active_quiz['subject']}")
            st.write(f"**Topic:** {active_quiz['topic']}")
            st.write(f"**Question:** {active_quiz['question']}")
            st.write(f"**Expected answer:** {active_quiz['expected_answer']}")
            st.write(f"**Created at:** {active_quiz['created_at']}")
        else:
            st.info("No active quiz question currently saved.")

        st.markdown("### Saved Learning Notes")
        notes = get_previous_notes(memory_student_name, limit=50)

        if not notes:
            st.info("No learning notes saved yet.")
        else:
            for note in notes:
                with st.expander(f"{note['created_at']} — {note['subject']}: {note['topic']}"):
                    st.write(note["note"])

        st.markdown("### Quiz Attempts")
        quiz_attempts = get_quiz_attempts(memory_student_name, limit=50)

        if not quiz_attempts:
            st.info("No quiz attempts recorded yet.")
        else:
            for item in quiz_attempts:
                with st.expander(
                    f"{item['created_at']} — {item['subject']}: {item['topic']} — {item['was_correct']}"
                ):
                    st.write(f"**Question:** {item['question']}")
                    st.write(f"**Student answer:** {item['student_answer']}")
                    st.write(f"**Correct answer:** {item['correct_answer']}")
                    st.write(f"**Feedback:** {item['feedback']}")

        st.markdown("### Recorded Struggle Areas")
        struggles = get_struggles(memory_student_name)

        if not struggles:
            st.info("No struggle areas recorded yet.")
        else:
            for item in struggles:
                with st.expander(f"{item['created_at']} — {item['subject']}: {item['topic']}"):
                    st.write(f"**Area:** {item['struggle_area']}")
                    st.write(f"**Evidence:** {item['evidence']}")
                    st.write(f"**Suggested support:** {item['suggested_support']}")
